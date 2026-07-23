from pathlib import Path
from html import escape
from zipfile import ZIP_DEFLATED, ZipFile


def _gerar_docx_sem_dependencia(destino):
    """Cria um DOCX OpenXML válido quando python-docx não está instalado."""
    secoes = [
        ('Funcionamento das APIs na Aplicação AgroVision', 'title'),
        ('Estudo funcional — Fazenda Cristalina, Malanje', 'subtitle'),
        ('1. Objetivo', 'h1'),
        ('Explicar como as APIs e o motor interno transformam dados da Fazenda Cristalina em meteorologia, risco, orientação, alertas e apoio ao chatbot.', 'p'),
        ('2. Cenário de validação', 'h1'),
        ('A Fazenda Cristalina está localizada em Malanje, Malanje, Angola e possui três talhões: feijão com 24,5 hectares, mandioca com 38 hectares e soja com 30 hectares.', 'p'),
        ('3. Fluxo técnico', 'h1'),
        ('1. O utilizador informa talhão, cultura e localização.', 'p'),
        ('2. O backend valida perfil, propriedade e qualidade dos dados.', 'p'),
        ('3. O AgroVision chama o provedor meteorológico ou o Gemini sem expor a chave.', 'p'),
        ('4. A resposta JSON é normalizada e combinada com solo, cultura, pragas e histórico.', 'p'),
        ('5. O motor calcula risco e orientação preliminar.', 'p'),
        ('6. O consultor revê e emite a recomendação oficial.', 'p'),
        ('4. APIs utilizadas', 'h1'),
        ('Open-Meteo e OpenWeatherMap: temperatura, humidade, vento, precipitação e previsão.', 'p'),
        ('Google Gemini: chatbot agrícola, explicações e pesquisa pública atual com Google Search quando disponível.', 'p'),
        ('NASA POWER: séries agroclimáticas históricas por coordenadas.', 'p'),
        ('SoilGrids: estimativas de propriedades do solo que exigem confirmação laboratorial.', 'p'),
        ('Sentinel-2/Copernicus: índices de vegetação e observação temporal dos talhões.', 'p'),
        ('FAOSTAT: contexto estatístico, sem substituir dados locais.', 'p'),
        ('5. Consultoria inteligente', 'h1'),
        ('A aplicação cruza clima, solo, cultura, alertas e pragas. Apresenta pontuação de risco, fatores, possíveis doenças e orientação. O agricultor envia a análise; o consultor usa Avaliar e responder e depois emite o PDF.', 'p'),
        ('6. Teste por utilizador', 'h1'),
        ('Messi consulta mercado e envia pedido de compra. Enock gere a fazenda e solicita consultoria. Ado avalia e emite recomendações. Garcia regista visitas e pragas. Priscila analisa indicadores. O administrador atribui equipas e configura APIs.', 'p'),
        ('7. Segurança e contingência', 'h1'),
        ('As chaves ficam no servidor. Os dados são limitados pelo perfil. Se a API falhar, o sistema informa a indisponibilidade e mantém regras e respostas locais. Possíveis doenças não constituem diagnóstico.', 'p'),
        ('8. Resultado', 'h1'),
        ('A chave Gemini respondeu em teste real. Os testes automatizados validaram permissões, fotografias, cultura livre, meteorologia, consultoria, envio, revisão, emissão e PDF.', 'p'),
        ('9. Roteiro de apresentação em 10 minutos', 'h1'),
        ('1 minuto para objetivo e perfis; 2 minutos para a fazenda; 3 minutos para API meteorológica e JSON; 2 minutos para risco e consultoria; 1 minuto para revisão; 1 minuto para Gemini, segurança e conclusão.', 'p'),
        ('10. Conclusão', 'h1'),
        ('O valor do AgroVision está no fluxo completo: dado externo, regra interna, explicação, revisão humana e entrega segura ao agricultor.', 'p'),
    ]
    paragrafos = []
    for texto, tipo in secoes:
        tamanho = '36' if tipo == 'title' else '30' if tipo == 'h1' else '24'
        negrito = '<w:b/>' if tipo in {'title', 'h1'} else ''
        alinhamento = '<w:jc w:val="center"/>' if tipo in {'title', 'subtitle'} else ''
        paragrafos.append(
            f'<w:p><w:pPr>{alinhamento}</w:pPr><w:r><w:rPr>{negrito}<w:sz w:val="{tamanho}"/></w:rPr>'
            f'<w:t xml:space="preserve">{escape(texto)}</w:t></w:r></w:p>'
        )
    linhas = [
        ('Talhão', 'Cultura', 'Área', 'Solo'),
        ('Talhão Cristal Feijão', 'Feijão', '24,5 ha', 'Franco-argiloso'),
        ('Talhão Raiz Cristalina', 'Mandioca', '38 ha', 'Arenoso bem drenado'),
        ('Talhão Horizonte Soja', 'Soja', '30 ha', 'Franco com matéria orgânica'),
    ]
    tabela = '<w:tbl>' + ''.join(
        '<w:tr>' + ''.join(
            f'<w:tc><w:p><w:r><w:t>{escape(valor)}</w:t></w:r></w:p></w:tc>' for valor in linha
        ) + '</w:tr>' for linha in linhas
    ) + '</w:tbl>'
    paragrafos.insert(5, tabela)
    documento = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>' + ''.join(paragrafos) + '<w:sectPr/></w:body></w:document>'
    )
    tipos = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '</Relationships>'
    )
    with ZipFile(destino, 'w', ZIP_DEFLATED) as arquivo:
        arquivo.writestr('[Content_Types].xml', tipos)
        arquivo.writestr('_rels/.rels', rels)
        arquivo.writestr('word/document.xml', documento)
    return destino


def gerar_relatorio(base_dir):
    """Gera o relatório técnico das APIs usando o cenário Fazenda Cristalina."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor
    except ModuleNotFoundError:
        base_dir = Path(base_dir)
        destino_dir = base_dir / 'relatorios'
        destino_dir.mkdir(parents=True, exist_ok=True)
        return _gerar_docx_sem_dependencia(
            destino_dir / 'Relatorio_Funcionamento_APIs_AgroVision_Fazenda_Cristalina.docx'
        )

    base_dir = Path(base_dir)
    destino_dir = base_dir / 'relatorios'
    destino_dir.mkdir(parents=True, exist_ok=True)
    destino = destino_dir / 'Relatorio_Funcionamento_APIs_AgroVision_Fazenda_Cristalina.docx'

    doc = Document()
    secao = doc.sections[0]
    secao.top_margin = Cm(2)
    secao.bottom_margin = Cm(2)
    secao.left_margin = Cm(2.2)
    secao.right_margin = Cm(2.2)
    estilos = doc.styles
    estilos['Normal'].font.name = 'Aptos'
    estilos['Normal'].font.size = Pt(11)
    for nome in ('Title', 'Heading 1', 'Heading 2'):
        estilos[nome].font.name = 'Aptos Display'
        estilos[nome].font.color.rgb = RGBColor(31, 111, 53)

    titulo = doc.add_heading('Funcionamento das APIs na Aplicação AgroVision', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_paragraph('Estudo funcional — Fazenda Cristalina, Malanje')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Data da validação: 23 de julho de 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER

    capa = Path(r'C:\Users\jason\AppData\Local\Temp\codex-clipboard-5a4001e3-2aba-420f-8eda-4a383aafdf89.jpg')
    if capa.exists():
        doc.add_picture(str(capa), width=Cm(13.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading('1. Objetivo', level=1)
    doc.add_paragraph(
        'Este relatório explica como as APIs e o motor interno do AgroVision transformam dados da '
        'Fazenda Cristalina em meteorologia, análise de risco, orientação agrícola, alertas e apoio '
        'ao chatbot. O cenário utiliza três culturas: feijão, mandioca e soja.'
    )

    doc.add_heading('2. Cenário de validação', level=1)
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    for celula, texto in zip(tabela.rows[0].cells, ['Talhão', 'Cultura', 'Área', 'Solo']):
        celula.text = texto
    dados = [
        ('Talhão Cristal Feijão', 'Feijão', '24,5 ha', 'Franco-argiloso'),
        ('Talhão Raiz Cristalina', 'Mandioca', '38 ha', 'Arenoso bem drenado'),
        ('Talhão Horizonte Soja', 'Soja', '30 ha', 'Franco com matéria orgânica'),
    ]
    for linha in dados:
        celulas = tabela.add_row().cells
        for celula, texto in zip(celulas, linha):
            celula.text = texto
    doc.add_paragraph(
        'Localização padronizada: Malanje, Malanje, Angola. A propriedade é atribuída ao agricultor '
        'Enock, ao consultor Ado, ao técnico Garcia e à analista Priscila. O cliente Messi consulta '
        'somente ofertas comerciais publicadas e autorizadas.'
    )

    doc.add_heading('3. Arquitetura de integração', level=1)
    etapas = [
        ('Entrada', 'O utilizador escolhe talhão, cultura e localização ou faz uma pergunta no chatbot.'),
        ('Validação', 'O backend confirma o perfil, a propriedade autorizada e a qualidade mínima dos dados.'),
        ('Consulta externa', 'O AgroVision chama o provedor meteorológico ou o Gemini sem expor a chave no navegador.'),
        ('Tratamento', 'A resposta JSON é convertida em unidades compreensíveis e combinada com solo, cultura, pragas e histórico.'),
        ('Decisão', 'Regras internas calculam risco, possíveis condições e prioridade; o resultado não é diagnóstico definitivo.'),
        ('Revisão humana', 'O agricultor envia a análise ao consultor, que revê e emite a recomendação oficial.'),
        ('Contingência', 'Se uma API estiver indisponível, o sistema informa a falha e mantém respostas e regras locais.'),
    ]
    for titulo_etapa, explicacao in etapas:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{titulo_etapa}: ').bold = True
        p.add_run(explicacao)

    doc.add_heading('4. APIs utilizadas', level=1)
    itens_api = [
        ('Open-Meteo / OpenWeatherMap', 'Obtêm temperatura, humidade, vento, precipitação, condição atual e previsão. A configuração decide o provedor ativo.'),
        ('Google Gemini', 'Apoia o chatbot agrícola e a redação explicativa. Com Google Search habilitado, pode fundamentar respostas em informação pública atual. O sistema envia apenas contexto limitado ao perfil.'),
        ('NASA POWER', 'Fonte prevista para séries agroclimáticas históricas e comparação de chuva, radiação e temperatura por coordenadas.'),
        ('SoilGrids', 'Fonte prevista para propriedades estimadas do solo. Os valores devem ser identificados como estimativas e confirmados por análise laboratorial.'),
        ('Sentinel-2/Copernicus', 'Fonte prevista para índices de vegetação e observação do talhão. Exige coordenadas, processamento de imagem e interpretação temporal.'),
        ('FAOSTAT', 'Fonte de contexto estatístico; não substitui observações da fazenda nem recomendações agronómicas locais.'),
    ]
    for nome, texto in itens_api:
        p = doc.add_paragraph()
        p.add_run(nome + ': ').bold = True
        p.add_run(texto)

    doc.add_heading('5. Consultoria inteligente no agricultor', level=1)
    doc.add_paragraph(
        'A consulta meteorológica usa a localização da propriedade. O motor cruza o resultado com '
        'cultura, solo, alertas e ocorrências de pragas. A saída contém risco de zero a cem, fatores '
        'considerados, possíveis doenças a observar e orientação preliminar. Ao clicar em Enviar '
        'análise ao consultor, é criado um rascunho. O consultor usa Avaliar e responder, corrige o '
        'texto e só então emite a recomendação e o PDF ao agricultor.'
    )

    doc.add_heading('6. Teste por utilizador', level=1)
    testes = [
        ('Messi — cliente', 'Consulta o mercado, fotografias e colheitas publicadas; envia pedido de compra, sem editar dados técnicos.'),
        ('Enock — agricultor', 'Cria a Fazenda Cristalina, talhões, culturas livres, fotografias e colheitas; executa a consultoria e envia ao consultor.'),
        ('Ado — consultor', 'Recebe rascunhos, avalia solo e clima, corrige prioridade e texto, emite a recomendação oficial.'),
        ('Garcia — técnico', 'Regista visitas, fotografias, pragas, sinais e ações observadas em campo.'),
        ('Priscila — analista', 'Consulta indicadores, clima, culturas, produção e históricos, mantendo acesso de leitura.'),
        ('Administrador', 'Atribui a equipa, configura APIs, aprova perfis, contactos comerciais e administra todos os registos.'),
    ]
    for perfil, validacao in testes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(perfil + ': ').bold = True
        p.add_run(validacao)

    doc.add_heading('7. Segurança e qualidade', level=1)
    doc.add_paragraph(
        'As chaves ficam no servidor e não são entregues ao navegador. Consultas exigem autenticação '
        'e são limitadas pelo perfil e pelas propriedades atribuídas. O chatbot não recebe dados de '
        'outros utilizadores. Recomendações automáticas são preliminares, possíveis doenças não são '
        'diagnósticos e doses de produtos exigem validação profissional e respeito ao rótulo.'
    )

    doc.add_heading('8. Resultado da validação', level=1)
    doc.add_paragraph(
        'A chave Gemini foi validada com resposta real. A suíte automatizada confirmou os painéis, '
        'permissões, cultura em texto livre, fotografias, meteorologia simulada, consultoria, envio ao '
        'consultor, resposta, emissão e PDF. O funcionamento externo continua sujeito à internet, '
        'quota do provedor e qualidade da localização informada.'
    )

    doc.add_heading('9. Roteiro de demonstração em 10 minutos', level=1)
    roteiro = [
        '1 minuto: objetivo do AgroVision e perfis.',
        '2 minutos: Fazenda Cristalina, talhões, culturas e fotografias.',
        '3 minutos: chamada meteorológica e explicação do JSON tratado.',
        '2 minutos: cálculo do risco e envio da análise ao consultor.',
        '1 minuto: avaliação e emissão da recomendação pelo consultor.',
        '1 minuto: chatbot Gemini, segurança, fallback e conclusão.',
    ]
    for item in roteiro:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('10. Conclusão', level=1)
    doc.add_paragraph(
        'As APIs são suficientes para demonstrar integração real, tratamento de dados, inteligência '
        'aplicada e colaboração entre utilizadores. A força do projeto está no fluxo completo: dado '
        'externo, regra interna, explicação, revisão humana e entrega segura ao agricultor.'
    )
    doc.save(destino)
    return destino


if __name__ == '__main__':
    gerar_relatorio(Path(__file__).resolve().parents[1])
