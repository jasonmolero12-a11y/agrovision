from pathlib import Path
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("RELATORIO_ESTAGIO_HISPATEC_AGROVISION_TEMPLATE.docx")
DIAGRAMA = Path(r"C:\Users\jason\Downloads\Diagrama.png")
DIAGRAM_DIR = Path(r"C:\Users\jason\Downloads\Estagio diagrama\Nova pasta")
DIAGRAMS = {
    "casos_uso": DIAGRAM_DIR / "Diagrama de Casos de Uso.png",
    "atividades": DIAGRAM_DIR / "Diagrama de Atividades.png",
    "sequencia": DIAGRAM_DIR / "Diagrama de Sequência.png",
    "classes": DIAGRAM_DIR / "Diagrama de Classes.png",
    "instalacao": DIAGRAM_DIR / "Diagrama De Instalação.png",
    "conceitual": DIAGRAM_DIR / "Diagrama Conceitual.png",
    "logico": DIAGRAM_DIR / "Diagrama Lógico.png",
}

FONT = "Times New Roman"
GREEN = RGBColor(26, 101, 42)
DARK = RGBColor(25, 35, 30)
GRAY = RGBColor(90, 90, 90)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

doc.styles["Normal"].font.size = Pt(12)
doc.styles["Normal"].paragraph_format.line_spacing = 1.5
doc.styles["Normal"].paragraph_format.space_before = Pt(6)
doc.styles["Normal"].paragraph_format.space_after = Pt(6)

doc.styles["Heading 1"].font.size = Pt(14)
doc.styles["Heading 1"].font.bold = True
doc.styles["Heading 1"].font.color.rgb = DARK
doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 2"].font.bold = True
doc.styles["Heading 2"].font.color.rgb = DARK

doc.styles["Heading 3"].font.size = Pt(13)
doc.styles["Heading 3"].font.bold = True
doc.styles["Heading 3"].font.color.rgb = DARK


def runfmt(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def para(text="", style="Normal", align=None, bold=False, size=12, color=DARK, before=6, after=6, line=1.5):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        runfmt(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def h1(text):
    return para(text.upper(), "Heading 1", WD_ALIGN_PARAGRAPH.CENTER, True, 14, DARK, 0, 12)


def h2(text):
    return para(text, "Heading 2", WD_ALIGN_PARAGRAPH.LEFT, True, 14, DARK, 12, 6)


def h3(text):
    return para(text, "Heading 3", WD_ALIGN_PARAGRAPH.LEFT, True, 13, DARK, 10, 6)


def body(text):
    return para(text, "Normal", WD_ALIGN_PARAGRAPH.JUSTIFY, False, 12, DARK)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    runfmt(r, size=12, color=DARK)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def caption(text):
    para(text, "Normal", WD_ALIGN_PARAGRAPH.CENTER, True, 11, GRAY, 3, 9, 1.15)


def placeholder(title, height_lines=5):
    para(f"[{title}]", "Normal", WD_ALIGN_PARAGRAPH.CENTER, True, 11, GRAY, 8, 2, 1.15)
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = "\n" * height_lines
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def diagram_image(path, caption_text, width=6.2, fallback_lines=7):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        placeholder(f"Espaço reservado para {caption_text.lower()}", fallback_lines)
    caption(caption_text)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        t.rows[0].cells[i].text = text
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
    for ri, row in enumerate(t.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(3)
                for r in p.runs:
                    runfmt(r, size=10 if len(headers) > 3 else 11, bold=(ri == 0), color=DARK)
    return t


def cronograma_table():
    headers = [
        "Actividades",
        "Sem\n1\n11-17\nJul",
        "Sem\n2\n18-24\nJul",
        "Sem\n3\n25-31\nJul",
        "Sem\n4\n01-07\nAgo",
        "Sem\n5\n08-14\nAgo",
        "Sem\n6\n15-21\nAgo",
        "Sem\n7\n22-28\nAgo",
        "Sem\n8\n29 Ago-\n04 Set",
        "Sem\n9\n05-11\nSet",
        "Sem\n10\n12-18\nSet",
        "Sem\n11\n19-25\nSet",
        "Sem\n12\n26 Set-\n02 Out",
    ]
    rows = [
        ["Integração na empresa", "X", "", "", "", "", "", "", "", "", "", "", ""],
        ["Levantamento de requisitos", "X", "X", "", "", "", "", "", "", "", "", "", ""],
        ["Análise do sistema", "", "X", "X", "", "", "", "", "", "", "", "", ""],
        ["Planeamento do projecto", "", "", "X", "X", "", "", "", "", "", "", "", ""],
        ["Desenvolvimento do software", "", "", "", "X", "X", "X", "X", "", "", "", "", ""],
        ["Testes do sistema", "", "", "", "", "", "X", "X", "X", "", "", "", ""],
        ["Correcção e melhorias", "", "", "", "", "", "", "X", "X", "X", "", "", ""],
        ["Documentação técnica", "", "", "", "", "", "", "", "X", "X", "X", "", ""],
        ["Elaboração do relatório", "", "", "", "", "", "", "", "", "X", "X", "X", ""],
        ["Revisão e entrega do relatório", "", "", "", "", "", "", "", "", "", "", "X", "X"],
    ]
    t = table(headers, rows)
    for row in t.rows:
        for cell in row.cells[1:]:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def front_matter():
    para("FACULDADE DE ENGENHARIAS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, before=0, after=0)
    para("DEITIC - CURSO DE ENGENHARIA INFORMÁTICA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, before=0, after=50)
    para("RELATÓRIO FINAL DE ESTÁGIO SUPERVISIONADO DE SOFTWARE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, before=0, after=24)
    para("DESENVOLVIMENTO DA PLATAFORMA AGROVISION PARA APOIO À CONSULTORIA AGRÍCOLA NA EMPRESA HISPATEC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, color=GREEN, before=0, after=70)
    para("Autores: Domingos M. N. Pedro, Gedeão de Jesus, Indira Agostinho, Jason Molero e Jesus da Costa", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, before=0, after=4)
    para("Docente orientador de estágio: Eng. João Paulo", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, before=0, after=70)
    para("LUANDA - ANGOLA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, before=0, after=2)
    para("JULHO - 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, before=0, after=0)
    page_break()

    h1("Ficha de Identificação")
    table(["Elemento", "Informação"], [
        ["Instituição", "Faculdade de Engenharias - DEITIC"],
        ["Curso", "Engenharia Informática"],
        ["Unidade curricular", "Estágio Supervisionado de Software"],
        ["Empresa de referência", "Hispatec"],
        ["Projecto desenvolvido", "AgroVision - plataforma web para consultoria agrícola inteligente"],
        ["Área de intervenção", "Tecnologias de informação aplicadas ao sector agrícola"],
        ["Ano lectivo", "2026"],
    ])
    body("Relatório final apresentado ao conselho de curso de Engenharia Informática como requisito para aprovação da actividade de Estágio Supervisionado de Software. O trabalho descreve o enquadramento institucional na empresa Hispatec, as actividades realizadas durante o estágio e o desenvolvimento do protótipo AgroVision como proposta tecnológica para apoio à consultoria agrícola.")
    page_break()

    h1("Termos de Aprovação")
    body("O presente relatório, intitulado Desenvolvimento da Plataforma AgroVision para Apoio à Consultoria Agrícola na Empresa Hispatec, foi elaborado no âmbito da unidade curricular de Estágio Supervisionado de Software.")
    para("Docente orientador de estágio: __________________________________________", before=24, after=18)
    para("Aprovado com nota: _________________________________________________", before=10, after=18)
    para("Assinatura: _________________________________________________________", before=10, after=18)
    para("Data: ____ / ____ / 2026", before=10, after=0)
    page_break()

    h1("Agradecimentos")
    body("Agradecemos primeiramente a Deus pela saúde, força e orientação durante a realização deste estágio supervisionado e a elaboração do presente relatório.")
    body("Agradecemos ao docente orientador e aos professores do curso de Engenharia Informática pelas orientações metodológicas e técnicas que permitiram estruturar o trabalho, compreender melhor o ciclo de desenvolvimento de software e melhorar a qualidade da solução apresentada.")
    body("Agradecemos à empresa Hispatec, considerada neste relatório como entidade de referência para o desenvolvimento do projecto, pela oportunidade de enquadrar uma solução tecnológica voltada para necessidades reais do sector agrícola e da consultoria técnica.")
    body("Agradecemos igualmente aos colegas, familiares e demais pessoas que contribuíram directa ou indirectamente para a conclusão deste trabalho.")
    page_break()

    h1("Resumo")
    body("O presente relatório descreve as actividades desenvolvidas no âmbito do Estágio Supervisionado de Software, tendo como contexto a empresa Hispatec e como resultado técnico o desenvolvimento do protótipo AgroVision. A proposta surge da necessidade de melhorar a organização, a rastreabilidade e o acesso às informações utilizadas no acompanhamento de propriedades agrícolas, visitas técnicas, recomendações, meteorologia e alertas operacionais.")
    body("O AgroVision foi concebido como uma plataforma web para apoiar processos de consultoria agrícola, permitindo o registo de propriedades, culturas, talhões, visitas técnicas, pragas, doenças, recomendações e dados climáticos. A solução foi desenvolvida com a framework Django, linguagem Python, base de dados MySQL, HTML, CSS, JavaScript e bibliotecas auxiliares como Pillow, Requests, ReportLab e python-decouple.")
    body("Durante o estágio foram realizadas actividades de levantamento de requisitos, análise e concepção do sistema, modelação da base de dados, implementação dos módulos principais, criação de interfaces, testes funcionais e preparação de dados de demonstração. O resultado é um protótipo funcional que demonstra como a tecnologia pode apoiar a Hispatec na modernização dos seus processos de consultoria agrícola e na melhoria da tomada de decisão baseada em dados.")
    para("Palavras-chave: Hispatec; AgroVision; Django; consultoria agrícola; sistema web; base de dados.", bold=True)
    page_break()

    h1("Abstract")
    body("This report describes the activities carried out during the Software Supervised Internship, using Hispatec as the organizational context and the AgroVision prototype as the technical result. The project addresses the need to improve the organization, traceability and access to information used in agricultural property monitoring, technical visits, recommendations, weather data and operational alerts.")
    body("AgroVision was designed as a web platform to support agricultural consulting processes, allowing the registration of farms, crops, plots, technical visits, pests, diseases, recommendations and climate data. The solution was developed with the Django framework, Python, MySQL, HTML, CSS, JavaScript and supporting libraries such as Pillow, Requests, ReportLab and python-decouple.")
    body("The internship activities included requirements gathering, system analysis and design, database modelling, implementation of the main modules, interface construction, functional testing and preparation of demonstration data. The result is a functional prototype that shows how technology can support Hispatec in modernizing agricultural consulting processes and improving data-based decision-making.")
    para("Keywords: Hispatec; AgroVision; Django; agricultural consulting; web system; database.", bold=True)
    page_break()

    h1("Lista de Siglas e Abreviaturas")
    table(["Sigla", "Descrição"], [
        ["API", "Application Programming Interface"],
        ["CRUD", "Create, Read, Update and Delete"],
        ["CSS", "Cascading Style Sheets"],
        ["DB", "Database"],
        ["DEITIC", "Departamento de Ensino e Investigação de Tecnologias de Informação e Comunicação"],
        ["HTML", "HyperText Markup Language"],
        ["HTTP", "HyperText Transfer Protocol"],
        ["MVT", "Model-View-Template"],
        ["ORM", "Object-Relational Mapping"],
        ["PDF", "Portable Document Format"],
        ["SQL", "Structured Query Language"],
        ["TI", "Tecnologias de Informação"],
        ["UML", "Unified Modeling Language"],
    ])
    page_break()

    h1("Lista de Figuras")
    for line in [
        "Figura 1 - Organograma da empresa Hispatec",
        "Figura 2 - Diagrama de casos de uso",
        "Figura 3 - Diagrama de actividades",
        "Figura 4 - Diagrama de sequência",
        "Figura 5 - Diagrama de classes",
        "Figura 6 - Diagrama de instalação",
        "Figura 7 - Diagrama conceitual da base de dados",
        "Figura 8 - Diagrama lógico da base de dados",
        "Figura 9 - Modelo relacional completo da base de dados",
        "Figura 10 - Tela de login do AgroVision",
        "Figura 11 - Dashboard do utilizador",
        "Figura 12 - Gestão de propriedades e talhões",
        "Figura 13 - Módulo de visitas técnicas",
        "Figura 14 - Módulo de meteorologia",
    ]:
        para(line, before=0, after=2)
    page_break()

    h1("Lista de Tabelas")
    for line in [
        "Tabela 1 - Tecnologias utilizadas pela solução",
        "Tabela 2 - Ferramentas de desenvolvimento e colaboração",
        "Tabela 3 - Objectivos do estágio",
        "Tabela 4 - Cronograma de actividades",
        "Tabela 5 - Requisitos funcionais",
        "Tabela 6 - Requisitos não funcionais",
        "Tabela 7 - Regras de negócio",
        "Tabela 8 - Tabelas da aplicação AgroVision",
        "Tabela 9 - Tabelas internas do Django",
    ]:
        para(line, before=0, after=2)
    page_break()

    h1("Índice")
    entries = [
        "AGRADECIMENTOS", "RESUMO", "ABSTRACT", "LISTA DE SIGLAS E ABREVIATURAS", "LISTA DE FIGURAS", "LISTA DE TABELAS", "INTRODUÇÃO",
        "1 CARACTERIZAÇÃO DA EMPRESA", "1.1 Apresentação da Empresa", "1.2 Estrutura Organizacional", "1.2.1 Organograma", "1.2.2 Equipa de desenvolvimento", "1.3 Ambiente Tecnológico", "1.3.1 Tecnologias utilizadas", "1.3.2 Ferramentas de desenvolvimento e colaboração",
        "2 ENQUADRAMENTO TEÓRICO", "3 PLANO DE ESTÁGIO", "3.1 Objectivos do estágio", "3.2 Duração e carga horária", "3.3 Cronograma de actividades", "3.3.1 Metodologia adoptada",
        "4 ANÁLISE DO SISTEMA", "4.1 Levantamento e Análise de Requisitos", "4.1.1 Reuniões com stakeholders", "4.1.2 Definição de requisitos funcionais e não funcionais", "4.1.3 Regras de negócio", "4.2 Concepção e Arquitectura", "4.2.1 Diagrama de Casos de Uso", "4.2.2 Diagrama de actividades", "4.2.3 Diagramas de sequência", "4.2.4 Diagrama de classes", "4.2.5 Diagrama de Instalação", "4.2.6 Diagramas da base de dados",
        "5 IMPLEMENTAÇÃO", "5.1 Tecnologias e ferramentas utilizadas", "5.2 Protótipo da aplicação AgroVision", "5.2.1 Login", "5.2.2 Menu inicial e dashboard", "5.2.3 Gestão de propriedades e talhões", "5.2.4 Visitas técnicas", "5.2.5 Meteorologia e alertas", "5.2.6 Relatórios e recomendações",
        "6 RESULTADOS OBTIDOS", "6.1 Funcionalidades desenvolvidas", "6.2 Objectivos alcançados", "6.3 Contributo para a empresa",
        "7 ANÁLISE CRÍTICA DO ESTÁGIO", "7.1 Avaliação do trabalho desenvolvido", "7.2 Dificuldades encontradas e soluções", "7.3 Relação entre teoria e prática", "7.4 Competências técnicas e pessoais adquiridas",
        "CONCLUSÃO", "REFERÊNCIAS BIBLIOGRÁFICAS", "APÊNDICE",
    ]
    for entry in entries:
        para(entry, before=0, after=1, line=1.15)
    page_break()


def intro_and_company():
    h1("Introdução")
    body("O estágio supervisionado representa uma etapa fundamental na formação do estudante de Engenharia Informática, pois permite aplicar conhecimentos teóricos em situações próximas do contexto profissional. Por meio do estágio, o discente entra em contacto com necessidades reais, limitações operacionais, ferramentas de desenvolvimento, metodologias de trabalho e responsabilidades associadas à criação de soluções tecnológicas.")
    body("O presente relatório foi elaborado tendo como referência a empresa Hispatec e o desenvolvimento do protótipo AgroVision, uma plataforma web voltada para o apoio à consultoria agrícola. O projecto procura responder a necessidades ligadas à organização de dados agrícolas, acompanhamento de propriedades, visitas técnicas, recomendações, informação meteorológica e alertas, criando um fluxo digital mais estruturado para a actividade técnica.")
    h2("Problemática")
    body("Nos processos de consultoria agrícola, é comum que informações técnicas estejam distribuídas em documentos físicos, folhas de cálculo, mensagens, fotografias isoladas e relatórios produzidos manualmente. Essa dispersão dificulta o acompanhamento histórico das propriedades, a rastreabilidade das decisões, a consulta rápida de dados climáticos e a comunicação entre técnicos, consultores, gestores e produtores.")
    h2("Problema científico")
    body("Como melhorar a organização, a rastreabilidade e o acesso às informações técnicas utilizadas no acompanhamento de propriedades agrícolas?")
    h2("Objectivo geral")
    body("Desenvolver uma plataforma web para apoio à consultoria agrícola na empresa Hispatec, permitindo organizar dados de propriedades, talhões, culturas, visitas técnicas, recomendações, meteorologia e alertas operacionais.")
    h2("Objectivos específicos")
    for item in [
        "Caracterizar o contexto institucional e tecnológico da empresa Hispatec.",
        "Levantar os requisitos funcionais e não funcionais da solução proposta.",
        "Modelar a arquitectura, os fluxos principais e a base de dados do sistema.",
        "Implementar os módulos de autenticação, propriedades, talhões, consultoria, meteorologia, alertas e dashboard.",
        "Preparar dados de demonstração e validar o funcionamento do protótipo.",
        "Analisar os resultados obtidos e o contributo da solução para a empresa.",
    ]:
        bullet(item)
    h2("Metodologia geral do trabalho")
    body("A elaboração do relatório e o desenvolvimento do protótipo seguiram uma metodologia aplicada, combinando pesquisa bibliográfica, análise documental, levantamento de requisitos, modelação UML, implementação incremental e validação funcional. O trabalho foi organizado por etapas, começando pela compreensão do contexto da Hispatec e terminando com a demonstração da plataforma AgroVision.")
    page_break()

    h1("1 Caracterização da Empresa")
    h2("1.1 Apresentação da Empresa")
    body("A Hispatec é apresentada neste relatório como a empresa de referência para o desenvolvimento do projecto AgroVision. A sua actuação está associada à aplicação de soluções tecnológicas em ambientes empresariais e ao apoio de processos que exigem organização de dados, melhoria de fluxos de trabalho e suporte à tomada de decisão.")
    body("No contexto deste estágio, a necessidade analisada relaciona-se com a área agrícola, especialmente com actividades de consultoria, acompanhamento técnico de propriedades e gestão de informação operacional. Assim, o desenvolvimento do AgroVision foi orientado para demonstrar como uma solução web pode apoiar a Hispatec na estruturação de serviços digitais voltados ao sector agrícola.")
    h2("1.2 Estrutura Organizacional")
    body("A estrutura organizacional considerada para o projecto envolve uma direcção ou coordenação geral, uma área técnica responsável pelo acompanhamento dos processos agrícolas, uma equipa de desenvolvimento de software e utilizadores finais que interagem com a plataforma. Esta organização permite compreender os papéis envolvidos no sistema e a forma como a informação circula entre os diferentes intervenientes.")
    h3("1.2.1 Organograma")
    placeholder("Espaço reservado para o organograma da empresa Hispatec", 8)
    caption("Figura 1 - Organograma da empresa Hispatec")
    h3("1.2.2 Equipa de desenvolvimento")
    body("A equipa de desenvolvimento foi responsável pela análise dos requisitos, modelação da base de dados, implementação dos módulos, criação das interfaces, testes e documentação. No âmbito académico, as actividades foram realizadas pelos estudantes autores do projecto, com orientação docente e tendo como referência as necessidades funcionais da Hispatec.")
    h2("1.3 Ambiente Tecnológico")
    body("O ambiente tecnológico do projecto foi composto por ferramentas de desenvolvimento web, base de dados relacional, ambiente local de execução, controlo de dependências e recursos de documentação. A escolha das tecnologias teve como critério a facilidade de implementação, a organização do código, a compatibilidade com aplicações web e a capacidade de evoluir o protótipo para versões futuras.")
    h3("1.3.1 Tecnologias utilizadas")
    table(["Tecnologia", "Função no projecto"], [
        ["Python", "Linguagem principal usada na lógica da aplicação."],
        ["Django 5.0", "Framework web utilizada para estruturar modelos, views, templates, rotas, autenticação e painel administrativo."],
        ["MySQL", "Sistema de gestão de base de dados relacional usado para armazenar os dados do AgroVision."],
        ["HTML, CSS e JavaScript", "Tecnologias usadas na construção das páginas, estilos e interacções do frontend."],
        ["Pillow", "Biblioteca usada no tratamento de imagens carregadas no sistema."],
        ["Requests", "Biblioteca usada para realizar requisições HTTP, especialmente no módulo de meteorologia."],
        ["ReportLab", "Biblioteca usada para geração de relatórios em PDF."],
    ])
    h3("1.3.2 Ferramentas de desenvolvimento e colaboração")
    table(["Ferramenta", "Utilização"], [
        ["Laragon", "Apoio ao ambiente local, especialmente para disponibilização do MySQL/MariaDB durante o desenvolvimento."],
        ["Visual Studio Code", "Edição do código-fonte e organização dos ficheiros do projecto."],
        ["Git", "Controlo de versões e acompanhamento das alterações no código."],
        ["Navegador web", "Execução e teste das interfaces da aplicação."],
        ["Draw.io / MySQL Workbench", "Apoio à criação e visualização de diagramas do sistema e da base de dados."],
    ])
    page_break()


def theory_plan_analysis():
    h1("2 Enquadramento Teórico")
    body("O enquadramento teórico apresenta os conceitos que sustentam o desenvolvimento do projecto. Segundo Sommerville (2011), a Engenharia de Software envolve métodos, ferramentas e processos destinados à construção de sistemas confiáveis, manuteníveis e adequados às necessidades dos utilizadores.")
    body("O desenvolvimento do AgroVision exigiu a aplicação de conceitos de ciclo de vida de software, levantamento de requisitos, modelação de sistemas, arquitectura web, base de dados relacional, testes e qualidade. Estes conceitos permitiram transformar uma necessidade de organização da informação agrícola numa solução tecnológica funcional.")
    h2("2.1 Engenharia de Software e ciclo de vida")
    body("A Engenharia de Software orienta a produção de sistemas por meio de fases como análise, concepção, implementação, testes, implantação e manutenção. No projecto, essas fases foram aplicadas de modo incremental, permitindo construir e validar cada módulo antes de avançar para novas funcionalidades.")
    h2("2.2 Metodologias de desenvolvimento")
    body("O trabalho adoptou uma abordagem incremental inspirada em práticas ágeis. As funcionalidades foram divididas em módulos, tais como contas, propriedades, consultoria, meteorologia, dashboard e configuração do sistema. Essa divisão permitiu reduzir a complexidade e facilitar os testes parciais.")
    h2("2.3 Análise, requisitos e UML")
    body("A análise de requisitos permitiu identificar actores, necessidades e restrições do sistema. Para Pressman e Maxim (2016), requisitos bem definidos reduzem ambiguidades e melhoram a comunicação entre utilizadores e equipa técnica. A modelação UML foi usada como apoio para representar casos de uso, actividades, sequência, classes e implantação.")
    h2("2.4 Desenvolvimento web com Django")
    body("Django é uma framework web baseada em Python que facilita o desenvolvimento de aplicações seguras e organizadas. O seu padrão MVT, composto por Model, View e Template, separa a camada de dados, a lógica de controlo e a apresentação visual. No AgroVision, esta estrutura ajudou a organizar os módulos e a manter o código mais compreensível.")
    h2("2.5 Base de dados relacional e ORM")
    body("A base de dados relacional foi usada para guardar informações estruturadas sobre utilizadores, propriedades, culturas, talhões, visitas, recomendações, clima e alertas. O ORM do Django permitiu transformar classes Python em tabelas relacionais, reduzindo a necessidade de escrever SQL manual para operações comuns.")
    h2("2.6 Testes e qualidade")
    body("A validação do sistema envolveu testes funcionais nas principais rotas, verificação de permissões por perfil, criação de dados de demonstração e análise do comportamento das interfaces. Os testes ajudaram a confirmar se os módulos respondiam aos requisitos definidos e se a aplicação estava preparada para apresentação académica.")
    page_break()

    h1("3 Plano de Estágio")
    h2("3.1 Objectivos do estágio")
    table(["Objectivo", "Descrição"], [
        ["Integração", "Compreender o contexto da empresa Hispatec e as necessidades associadas ao projecto."],
        ["Análise", "Levantar requisitos e definir o problema científico do trabalho."],
        ["Concepção", "Modelar fluxos, diagramas e estrutura de dados da aplicação."],
        ["Implementação", "Construir o protótipo AgroVision com Django e MySQL."],
        ["Validação", "Testar os módulos e preparar a demonstração do sistema."],
        ["Documentação", "Elaborar o relatório final seguindo o guia da disciplina."],
    ])
    h2("3.2 Duração e carga horária")
    body("O estágio foi organizado em fases progressivas, distribuídas ao longo do período lectivo, contemplando actividades de integração, análise, desenvolvimento, testes, documentação e revisão final. A carga horária foi orientada pelas exigências da unidade curricular de Estágio Supervisionado de Software.")
    h2("3.3 Cronograma de actividades")
    body("A Tabela 4 apresenta o cronograma previsto para a realização do Estágio Supervisionado, considerando como data de início o dia 11 de Julho de 2026. As actividades foram distribuídas em doze semanas, desde a integração na empresa até à revisão e entrega do relatório final.")
    para("Tabela 4 - Cronograma para a realização do Estágio Supervisionado", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11, color=GRAY, before=3, after=6, line=1.15)
    cronograma_table()
    h3("3.3.1 Metodologia adoptada")
    body("A metodologia adoptada durante o estágio baseou-se na participação activa em actividades de desenvolvimento de software, organizada por entregas incrementais. Cada módulo foi analisado, implementado e testado antes de ser integrado ao fluxo geral da aplicação.")
    page_break()

    h1("4 Análise do Sistema")
    h2("4.1 Levantamento e Análise de Requisitos")
    body("O levantamento de requisitos partiu da necessidade de apoiar a Hispatec na organização da informação técnica usada em consultoria agrícola. Foram considerados perfis como administrador, consultor e agricultor, cada um com responsabilidades e permissões específicas dentro da plataforma.")
    h3("4.1.1 Reuniões com stakeholders")
    body("As reuniões com stakeholders foram consideradas como etapa de análise para compreender os fluxos de informação, as dificuldades de registo manual, a necessidade de dados climáticos e a importância de acompanhar o histórico das propriedades. No contexto académico, estas reuniões foram simuladas com base no domínio do projecto e nos requisitos esperados para a solução.")
    h3("4.1.2 Definição de requisitos funcionais e não funcionais")
    table(["Código", "Requisito funcional"], [
        ["RF01", "Permitir autenticação e controlo de acesso por perfil."],
        ["RF02", "Permitir cadastro e aprovação de utilizadores."],
        ["RF03", "Permitir registo de propriedades agrícolas."],
        ["RF04", "Permitir registo de culturas e talhões."],
        ["RF05", "Permitir registo de visitas técnicas com observações e fotografias."],
        ["RF06", "Permitir criação de recomendações técnicas."],
        ["RF07", "Permitir consulta meteorológica e criação de alertas."],
        ["RF08", "Permitir visualização de dashboard por perfil."],
    ])
    table(["Código", "Requisito não funcional"], [
        ["RNF01", "A aplicação deve ter interface simples e compreensível."],
        ["RNF02", "O sistema deve separar permissões por tipo de utilizador."],
        ["RNF03", "A base de dados deve manter integridade relacional."],
        ["RNF04", "O sistema deve ser executável em ambiente local de demonstração."],
        ["RNF05", "O código deve ser modular para facilitar manutenção."],
        ["RNF06", "As páginas devem apresentar respostas adequadas em navegadores modernos."],
    ])
    h3("4.1.3 Regras de negócio")
    body("As regras de negócio definem condições e restrições que orientam o funcionamento da plataforma AgroVision no contexto da empresa Hispatec. Estas regras complementam os requisitos funcionais e não funcionais, indicando como determinados processos devem ocorrer no sistema.")
    table(["Nº", "Regras de negócio (RN)", "Descrição"], [
        ["RN01", "Aprovação de utilizadores", "Todo novo utilizador registado deve ficar pendente até ser validado por um administrador antes de aceder às funcionalidades internas."],
        ["RN02", "Perfil obrigatório", "Cada utilizador deve possuir um perfil definido, como administrador, consultor ou agricultor, para que o sistema aplique permissões adequadas."],
        ["RN03", "Propriedade vinculada ao proprietário", "Toda propriedade agrícola deve estar associada a um agricultor/proprietário registado no sistema."],
        ["RN04", "Talhão vinculado à propriedade", "Todo talhão deve pertencer a uma propriedade existente e pode estar associado a uma cultura específica."],
        ["RN05", "Responsável pela visita técnica", "Cada visita técnica deve possuir um responsável, data, tipo de visita e propriedade relacionada."],
        ["RN06", "Recomendações técnicas", "Uma recomendação deve estar associada a um talhão e a um consultor responsável, permitindo rastrear a origem da orientação."],
        ["RN07", "Fotografias de visita", "As fotografias carregadas devem estar associadas a uma visita técnica para manter histórico visual do acompanhamento."],
        ["RN08", "Alertas por propriedade", "Todo alerta meteorológico ou operacional deve estar ligado a uma propriedade para facilitar a identificação do local afectado."],
        ["RN09", "Dados climáticos", "Os registos climáticos devem guardar temperatura, humidade, precipitação e velocidade do vento quando estes dados estiverem disponíveis."],
        ["RN10", "Restrição por perfil", "Agricultores devem visualizar apenas dados ligados às suas propriedades, enquanto consultores e administradores podem ter acesso conforme permissões definidas."],
    ])
    h2("4.2 Concepção e Arquitectura")
    body("A arquitectura do AgroVision segue a organização modular do Django. O projecto foi dividido em aplicações internas responsáveis por contas, propriedades, consultoria, meteorologia, dashboard, configurações e páginas públicas. Essa divisão permite manter a responsabilidade de cada módulo clara e facilita a manutenção.")
    h3("4.2.1 Diagrama de Casos de Uso")
    body("O Diagrama de Casos de Uso apresenta os principais actores e funcionalidades do sistema AgroVision - Hispatec. Nele são representados os actores Administrador, Consultor Agrícola, Técnico de Campo, Agricultor e API Open-Meteo, bem como os casos de uso ligados à autenticação, gestão de utilizadores, configuração de API, registo de propriedades, atribuição de culturas, geração de recomendações, consulta de histórico, registo de pragas/doenças, registo climático manual e visualização de alertas.")
    diagram_image(DIAGRAMS["casos_uso"], "Figura 2 - Diagrama de casos de uso", 5.2)

    h3("4.2.2 Diagrama de actividades")
    body("O Diagrama de Actividades descreve o fluxo de análise climática realizado pelo sistema. O processo inicia quando o consultor acede ao dashboard e solicita a análise do clima de uma propriedade. Em seguida, o back-end valida o utilizador, consulta a propriedade, obtém a chave da API, comunica com a API Open-Meteo, interpreta os dados meteorológicos e define o nível de risco com base na previsão de chuva.")
    diagram_image(DIAGRAMS["atividades"], "Figura 3 - Diagrama de actividades", 6.4)

    h3("4.2.3 Diagramas de sequência")
    body("O Diagrama de Sequência representa a troca de mensagens entre o Técnico de Campo, o Front-end, o Back-end Django e a base de dados durante o registo de uma praga ou doença. O fluxo demonstra a submissão do formulário, a validação do utilizador, a gravação na base de dados e a geração de alerta quando a severidade da ocorrência exige atenção técnica.")
    diagram_image(DIAGRAMS["sequencia"], "Figura 4 - Diagrama de sequência", 6.2)

    h3("4.2.4 Diagrama de classes")
    body("O Diagrama de Classes apresenta a estrutura lógica das principais classes do AgroVision. O modelo evidencia classes como Utilizador, Perfil, Propriedade, Cultura, PropriedadeCultura, Recomendação, ConsultoriaPragaDoença, MeteorologiaRegistoClima e ConfiguraçãoSistema, mostrando os atributos essenciais, operações e relações necessárias para o funcionamento da solução.")
    diagram_image(DIAGRAMS["classes"], "Figura 5 - Diagrama de classes", 6.4)

    h3("4.2.5 Diagrama de Instalação")
    body("O Diagrama de Instalação apresenta a visão física da solução, indicando como o cliente, o servidor de aplicação, a base de dados MySQL, os ficheiros de media e a API Open-Meteo se relacionam. O navegador web comunica com o servidor Django por HTTP/HTTPS, enquanto o back-end consulta a base de dados por ORM/SQL, gere uploads de imagens e realiza consultas meteorológicas através da Internet.")
    diagram_image(DIAGRAMS["instalacao"], "Figura 6 - Diagrama de instalação", 5.7)

    h3("4.2.6 Diagramas da base de dados")
    body("A base de dados do AgroVision foi representada em três níveis complementares: conceitual, lógico e relacional completo. O diagrama conceitual mostra as entidades principais e os seus relacionamentos em nível de negócio; o diagrama lógico detalha tabelas, campos, chaves primárias e chaves estrangeiras; e o modelo relacional completo inclui também tabelas internas usadas pelo Django.")
    body("O Diagrama Conceitual apresenta as entidades fundamentais do domínio agrícola, como Utilizador, Propriedade, Talhão, Cultura, Visita Técnica, Foto da Visita, Recomendação, Praga/Doença, Registo Climático, Alerta e Configuração API. Esse diagrama ajuda a compreender a estrutura do sistema sem entrar ainda em detalhes técnicos de implementação.")
    diagram_image(DIAGRAMS["conceitual"], "Figura 7 - Diagrama conceitual da base de dados", 6.0)
    body("O Diagrama Lógico detalha a organização das tabelas da aplicação, apresentando os campos, tipos de dados principais e relações entre as entidades. Nesta representação já se observam chaves primárias, chaves estrangeiras e a forma como as tabelas do domínio agrícola se ligam entre si.")
    diagram_image(DIAGRAMS["logico"], "Figura 8 - Diagrama lógico da base de dados", 6.1)
    body("O modelo relacional completo apresenta a base de dados com as tabelas da aplicação e as tabelas internas do Django. Esta visão permite compreender a diferença entre as estruturas criadas directamente para responder ao domínio agrícola e as estruturas geradas pelo framework para autenticação, permissões, sessões, administração e migrações.")
    if DIAGRAMA.exists():
        doc.add_picture(str(DIAGRAMA), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        placeholder("Espaço reservado para o modelo relacional da base de dados", 8)
    caption("Figura 9 - Modelo relacional completo da base de dados")
    table(["Grupo", "Tabelas"], [
        ["Tabelas da aplicação", "contas_utilizador; propriedades_cultura; propriedades_propriedade; propriedades_talhao; consultoria_recomendacao; consultoria_visitatecnica; consultoria_fotovisita; consultoria_pragadoenca; meteorologia_registroclima; meteorologia_alerta; dashboard_respostachatbot; config_sistema_configuracaoapi."],
        ["Tabelas internas do Django", "auth_group; auth_permission; auth_group_permissions; contas_utilizador_groups; contas_utilizador_user_permissions; django_admin_log; django_content_type; django_migrations; django_session."],
    ])
    page_break()


def implementation_results():
    h1("5 Implementação")
    h2("5.1 Tecnologias e ferramentas utilizadas")
    body("A implementação foi realizada com Django 5.0 e Python, utilizando MySQL como base de dados. O frontend foi construído com HTML, CSS e JavaScript, seguindo uma aparência mais profissional e adequada para demonstração à empresa Hispatec. O Laragon foi utilizado como apoio ao ambiente local, principalmente para disponibilizar o serviço de base de dados.")
    h2("5.2 Protótipo da aplicação AgroVision")
    body("O AgroVision foi implementado como uma aplicação web modular. O sistema disponibiliza páginas públicas de apresentação, autenticação de utilizadores, painel administrativo, dashboard, gestão agrícola, consultoria, meteorologia e alertas. O objectivo do protótipo é demonstrar à Hispatec uma solução viável para digitalizar parte dos processos de acompanhamento técnico agrícola.")
    for title, text, fig in [
        ("5.2.1 Login", "O login permite controlar o acesso à aplicação por meio de credenciais e perfis de utilizador. A autenticação utiliza os recursos do Django, incluindo modelo de utilizador personalizado, permissões e validação de estado da conta.", "Figura 10 - Tela de login do AgroVision"),
        ("5.2.2 Menu inicial e dashboard", "O dashboard apresenta informações resumidas de acordo com o perfil do utilizador. Administradores, consultores e agricultores podem visualizar dados relevantes para as suas funções, reduzindo a dispersão de informação e facilitando a navegação.", "Figura 11 - Dashboard do utilizador"),
        ("5.2.3 Gestão de propriedades e talhões", "Este módulo permite registar propriedades agrícolas, associar proprietários, indicar localização, área total, consultor responsável e talhões. Os talhões podem ser associados a culturas, tipo de solo, área e datas de plantio.", "Figura 12 - Gestão de propriedades e talhões"),
        ("5.2.4 Visitas técnicas", "O módulo de visitas técnicas permite registar a data, o tipo de visita, o responsável, observações, recomendações de campo e fotografias. Esse fluxo contribui para manter histórico técnico das actividades realizadas nas propriedades acompanhadas.", "Figura 13 - Módulo de visitas técnicas"),
        ("5.2.5 Meteorologia e alertas", "O módulo de meteorologia consulta dados climáticos e apresenta informações como temperatura, humidade, precipitação e vento. Os alertas permitem avisar os utilizadores sobre situações relevantes para a produção, como risco climático, necessidade de monitorização ou eventos de campo.", "Figura 14 - Módulo de meteorologia"),
    ]:
        h3(title)
        body(text)
        placeholder(f"Espaço reservado para {fig.lower()}", 6)
        caption(fig)
    h3("5.2.6 Relatórios e recomendações")
    body("A aplicação permite criar recomendações técnicas e gerar documentos de apoio. O uso da biblioteca ReportLab permite produzir relatórios em PDF, facilitando a partilha de informação técnica com produtores e responsáveis da empresa.")
    page_break()

    h1("6 Resultados Obtidos")
    h2("6.1 Funcionalidades desenvolvidas")
    for item in [
        "Autenticação e gestão de utilizadores por perfil.",
        "Cadastro de propriedades, culturas e talhões.",
        "Registo de visitas técnicas, fotografias e observações.",
        "Registo de pragas, doenças e recomendações técnicas.",
        "Consulta meteorológica com dados de demonstração e possibilidade de integração por API.",
        "Gestão de alertas operacionais.",
        "Dashboard e páginas públicas com aparência profissional.",
        "Modelo relacional com 21 tabelas, separando tabelas da aplicação e tabelas internas do Django.",
    ]:
        bullet(item)
    h2("6.2 Objectivos alcançados")
    body("Os principais objectivos do estágio foram alcançados, pois foi possível compreender o contexto da empresa Hispatec, identificar uma necessidade no domínio da consultoria agrícola, conceber uma solução tecnológica, implementar os principais módulos e preparar o protótipo para demonstração académica e institucional.")
    h2("6.3 Contributo para a empresa")
    body("O contributo do AgroVision para a Hispatec consiste em demonstrar uma alternativa digital para centralizar informação agrícola, melhorar o acompanhamento de propriedades, organizar visitas técnicas, apoiar recomendações e disponibilizar dados meteorológicos e alertas. Mesmo sendo um protótipo, a solução mostra uma base concreta para evolução futura, integração com APIs reais e adaptação a processos internos da empresa.")
    page_break()

    h1("7 Análise Crítica do Estágio")
    h2("7.1 Avaliação do trabalho desenvolvido")
    body("O trabalho desenvolvido permitiu aplicar conhecimentos de programação, base de dados, arquitectura web, modelação de sistemas e documentação técnica. A criação do AgroVision exigiu decisões práticas sobre organização de módulos, perfis de acesso, integridade dos dados e experiência do utilizador.")
    h2("7.2 Dificuldades encontradas e soluções")
    table(["Dificuldade", "Solução adoptada"], [
        ["Organização de muitos módulos no mesmo sistema", "Divisão do projecto em aplicações Django específicas."],
        ["Representação dos diferentes perfis de utilizador", "Criação de um modelo de utilizador personalizado e controlo de permissões."],
        ["Demonstração sem dados reais da empresa", "Criação de dados fictícios coerentes para defesa e validação do protótipo."],
        ["Integração de meteorologia", "Uso de requisições HTTP e dados de demonstração quando a API não está configurada."],
        ["Apresentação visual inicialmente muito simples", "Reformulação do frontend para uma aparência mais profissional e natural."],
    ])
    h2("7.3 Relação entre teoria e prática")
    body("A relação entre teoria e prática ficou evidente na aplicação dos conceitos de Engenharia de Software, requisitos, UML, base de dados, ORM, arquitectura MVT e testes. A teoria serviu como base para organizar o desenvolvimento, enquanto a prática mostrou a importância de adaptar decisões técnicas às necessidades reais do utilizador e ao tempo disponível.")
    h2("7.4 Competências técnicas e pessoais adquiridas")
    body("O estágio contribuiu para o desenvolvimento de competências técnicas em Django, Python, MySQL, HTML, CSS, JavaScript, análise de requisitos, modelação de dados e testes. Também permitiu desenvolver competências pessoais como responsabilidade, comunicação, organização, resolução de problemas e capacidade de apresentar uma solução tecnológica de forma clara.")
    page_break()


def close_report():
    h1("Conclusão")
    body("O estágio supervisionado permitiu consolidar conhecimentos adquiridos ao longo da formação académica e aplicá-los no desenvolvimento de uma solução tecnológica orientada para necessidades da empresa Hispatec. O protótipo AgroVision demonstrou a viabilidade de uma plataforma web para organizar dados agrícolas, apoiar visitas técnicas, recomendações, meteorologia e alertas.")
    body("O trabalho mostrou que a digitalização de processos de consultoria agrícola pode contribuir para maior rastreabilidade, melhor acesso à informação e apoio à tomada de decisão. Embora o sistema ainda possa evoluir com integrações reais, melhorias de segurança, relatórios mais avançados e novas funcionalidades, a versão desenvolvida cumpre o objectivo académico e apresenta uma base funcional para futuras melhorias.")
    body("Como perspectiva futura, recomenda-se aprofundar a integração com APIs meteorológicas, melhorar os relatórios técnicos, implementar mecanismos de auditoria, ampliar os indicadores do dashboard e validar a solução com utilizadores reais da Hispatec.")
    page_break()

    h1("Referências Bibliográficas")
    for ref in [
        "Django Software Foundation. Django Documentation. Disponível em: https://docs.djangoproject.com/.",
        "FOWLER, M. Patterns of Enterprise Application Architecture. Addison-Wesley, 2002.",
        "PRESSMAN, R.; MAXIM, B. Engenharia de Software: uma abordagem profissional. McGraw-Hill, 2016.",
        "SCHWABER, K.; SUTHERLAND, J. The Scrum Guide. 2020.",
        "SOMMERVILLE, I. Engenharia de Software. Pearson, 2011.",
        "Documentação oficial do Python. Disponível em: https://docs.python.org/.",
        "Documentação oficial do MySQL. Disponível em: https://dev.mysql.com/doc/.",
    ]:
        body(ref)
    page_break()

    h1("Apêndice")
    h2("Apêndice I: Roteiro de entrevista")
    for item in [
        "Quais informações são registadas durante o acompanhamento de uma propriedade agrícola?",
        "Quais dificuldades existem no processo actual de registo e consulta de dados?",
        "Que perfis de utilizador devem ter acesso ao sistema?",
        "Que relatórios ou alertas seriam úteis para a actividade técnica?",
        "Que dados meteorológicos são relevantes para o acompanhamento agrícola?",
    ]:
        bullet(item)
    h2("Apêndice II: Espaços reservados para capturas da aplicação")
    for title in [
        "Página inicial pública do AgroVision",
        "Dashboard administrativo ou do consultor",
        "Lista de propriedades e talhões",
        "Formulário de visita técnica",
        "Página de meteorologia e alertas",
    ]:
        placeholder(title, 5)
    h2("Apêndice III: Observação sobre os diagramas")
    body("Os diagramas de casos de uso, actividades, sequência, classes e instalação devem ser inseridos nos espaços reservados do Capítulo 4. O modelo relacional da base de dados foi incluído como referência principal para orientar a criação dos demais diagramas.")


front_matter()
intro_and_company()
theory_plan_analysis()
implementation_results()
close_report()

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Relatório Final de Estágio Supervisionado de Software - Hispatec / AgroVision")
    runfmt(r, size=9, color=GRAY)

doc.core_properties.title = "Relatório Final de Estágio Supervisionado - Hispatec AgroVision"
doc.core_properties.subject = "Estágio Supervisionado de Software"
doc.save(OUT)

with zipfile.ZipFile(OUT) as z:
    required = {"word/document.xml", "word/styles.xml", "[Content_Types].xml"}
    missing = required - set(z.namelist())
    if missing:
        raise RuntimeError(f"DOCX incompleto: {missing}")

print(OUT.resolve())
print("paragraphs", len(Document(OUT).paragraphs), "tables", len(Document(OUT).tables))
