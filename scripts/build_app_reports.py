from pathlib import Path
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_FUNC = Path("RELATORIO_FUNCIONAMENTO_AGROVISION.docx")
OUT_TECH = Path("RELATORIO_TECNICO_PROGRAMACAO_AGROVISION.docx")

FONT = "Arial"
GREEN = RGBColor(31, 112, 51)
DARK = RGBColor(37, 45, 40)
GRAY = RGBColor(95, 95, 95)
LIGHT = RGBColor(235, 244, 237)


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = GREEN
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)

    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 2"].font.color.rgb = DARK
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)

    add_title_page(doc, title, subtitle)
    return doc


def runfmt(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def p(doc, text="", style="Normal", align=None, bold=False, size=None, color=None):
    par = doc.add_paragraph(style=style)
    if text:
        r = par.add_run(text)
        runfmt(r, size=size, bold=bold, color=color or DARK)
    if align is not None:
        par.alignment = align
    return par


def h1(doc, text):
    return p(doc, text, "Heading 1")


def h2(doc, text):
    return p(doc, text, "Heading 2")


def bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    r = par.add_run(text)
    runfmt(r, size=11, color=DARK)
    par.paragraph_format.space_after = Pt(3)
    return par


def number(doc, text):
    par = doc.add_paragraph(style="List Number")
    r = par.add_run(text)
    runfmt(r, size=11, color=DARK)
    par.paragraph_format.space_after = Pt(3)
    return par


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        tbl.rows[0].cells[i].text = head
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for par in cell.paragraphs:
                par.paragraph_format.line_spacing = 1.05
                par.paragraph_format.space_after = Pt(2)
                for run in par.runs:
                    runfmt(run, size=9 if len(headers) > 3 else 10, bold=(ri == 0), color=DARK)
    return tbl


def add_title_page(doc, title, subtitle):
    p(doc, "AgroVision / Hispatec", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=GREEN)
    doc.add_paragraph()
    p(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20, color=GREEN)
    p(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=GRAY)
    doc.add_paragraph()
    p(doc, "Documento de apoio para defesa, uso e continuidade do projecto.", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=DARK)
    doc.add_paragraph()
    p(doc, "Data: Julho de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=GRAY)
    page_break(doc)


def footer(doc, text):
    for section in doc.sections:
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = f.add_run(text)
        runfmt(r, size=8, color=GRAY)


def build_functional_report():
    doc = setup_doc(
        "Relatório de Funcionamento da Aplicação AgroVision",
        "Como a plataforma funciona para utilizadores, defesa e demonstração",
    )

    h1(doc, "1. Visão Geral")
    p(doc, "O AgroVision é uma aplicação web desenvolvida para apoiar a consultoria agrícola no contexto da empresa Hispatec. A plataforma organiza dados sobre propriedades, talhões, culturas, visitas técnicas, pragas/doenças, recomendações, meteorologia e alertas. A versão actual deve ser apresentada como um protótipo funcional, preparado para demonstração académica e evolução futura.")
    p(doc, "A aplicação possui integração opcional com Google/Gemini API para apoiar o chatbot e a geração automática de recomendações. Quando a chave não está configurada, a internet falha ou a API não responde, o sistema continua a funcionar com respostas programadas e regras internas, evitando que a demonstração fique dependente da ligação externa.")

    h1(doc, "2. Objectivo da Aplicação")
    bullet(doc, "Centralizar informações agrícolas que antes poderiam estar dispersas em documentos, mensagens e folhas de cálculo.")
    bullet(doc, "Permitir acompanhamento de propriedades, culturas e talhões.")
    bullet(doc, "Registar visitas técnicas, fotografias, pragas/doenças e recomendações.")
    bullet(doc, "Consultar dados meteorológicos e criar alertas simples por regras.")
    bullet(doc, "Separar o acesso por perfil para proteger dados de cada utilizador.")

    h1(doc, "3. Perfis de Utilizador")
    table(doc, ["Perfil", "Função principal", "Acesso esperado"], [
        ["Administrador", "Gerir utilizadores, aprovar perfis, configurar API e consultar dados gerais.", "Acesso amplo ao sistema e ao painel administrativo."],
        ["Visitante", "Conta criada publicamente antes da aprovação.", "Só solicita perfil e aguarda validação."],
        ["Consultor Agrícola", "Acompanhar propriedades atribuídas e emitir recomendações.", "Recomendações, visitas, pragas e meteorologia das propriedades sob responsabilidade."],
        ["Técnico de Campo", "Registar visitas e ocorrências observadas em campo.", "Dados operacionais de campo, visitas e pragas/doenças."],
        ["Agricultor", "Consultar informações das suas propriedades.", "Propriedades próprias, recomendações, visitas e alertas associados."],
        ["Cliente", "Acompanhar relatórios e recomendações ligadas ao seu acesso.", "Relatórios, recomendações e histórico autorizado."],
        ["Analista", "Consultar indicadores e dados consolidados.", "Visão analítica e estatística."],
    ])

    h1(doc, "4. Fluxo de Entrada e Aprovação")
    number(doc, "O utilizador acede à página pública e cria uma conta.")
    number(doc, "O sistema atribui o perfil Visitante por segurança.")
    number(doc, "Depois do login, o utilizador solicita o perfil pretendido.")
    number(doc, "O utilizador preenche justificativa, validação profissional e anexa CV em PDF.")
    number(doc, "O administrador analisa o pedido no painel admin.")
    number(doc, "Se aprovado, o sistema muda o perfil para consultor, agricultor, técnico, cliente ou analista.")

    h1(doc, "5. Módulos da Aplicação")
    h2(doc, "5.1 Portal Público")
    p(doc, "Apresenta a página inicial, serviços, informações sobre o projecto e contacto. Serve como porta de entrada para visitantes e utilizadores que ainda não entraram no sistema.")
    h2(doc, "5.2 Contas e Perfil")
    p(doc, "Controla registo, login, logout, perfil, solicitação de acesso e lista de utilizadores. O login aceita email ou nome completo, conforme o backend personalizado da aplicação. O painel mostra aviso quando existem solicitações de perfil pendentes.")
    h2(doc, "5.3 Dashboard")
    p(doc, "Após o login, cada utilizador entra num dashboard adaptado ao seu perfil. O administrador vê totais gerais; o consultor vê recomendações e visitas; o agricultor vê dados ligados às suas propriedades; e o visitante vê o estado da solicitação.")
    h2(doc, "5.4 Propriedades, Culturas e Talhões")
    p(doc, "Permite registar propriedades agrícolas, culturas e talhões. As propriedades podem ter proprietário, consultor responsável, localização, latitude, longitude e área. Os talhões ficam ligados a uma propriedade e podem receber culturas.")
    h2(doc, "5.5 Consultoria")
    p(doc, "Inclui recomendações agronómicas, visitas técnicas, fotografias de visitas e registos de pragas/doenças. As recomendações podem ser exportadas em PDF. O sistema tenta gerar uma recomendação automática com Gemini quando a IA está configurada; se não estiver, usa o motor de regras interno. Em ambos os casos, o consultor deve rever antes de emitir.")
    h2(doc, "5.6 Meteorologia")
    p(doc, "Consulta previsão meteorológica através de API externa. O sistema suporta Open-Meteo e OpenWeatherMap. A pesquisa aceita cidades, províncias e municípios agrícolas de Angola, como Huambo, Cela, Caconda, Ganda, Matala e outros. Quando há propriedade associada, os dados podem ser guardados no histórico e podem gerar alertas simples, como calor extremo ou vento forte.")
    h2(doc, "5.7 Alertas")
    p(doc, "Os alertas avisam os utilizadores sobre situações relevantes ligadas a propriedades. Eles podem ser gerados por regras simples a partir dos dados climáticos ou por ocorrências de campo, como pragas/doenças com severidade alta ou crítica. O cabeçalho do painel apresenta o número de alertas não lidos.")
    h2(doc, "5.8 Configuração da API")
    p(doc, "Área reservada ao administrador. Permite escolher o provedor meteorológico, inserir chave quando necessário, configurar a chave Google/Gemini, indicar o modelo de IA, activar ou desligar a IA externa e testar a ligação meteorológica.")
    h2(doc, "5.9 Chatbot Interno")
    p(doc, "O chatbot trabalha em dois modos. Primeiro tenta usar Gemini quando a IA externa está activa e a chave está configurada. Se a API falhar, responde automaticamente com regras programadas e respostas editáveis guardadas no sistema. Em qualquer modo, o contexto enviado é limitado ao perfil do utilizador.")
    h2(doc, "5.10 Apoio por Voz")
    p(doc, "A aplicação possui apoio por voz no painel interno. Ao entrar no sistema, o utilizador ouve automaticamente uma mensagem de boas-vindas. Também pode ouvir a página actual, respostas do chatbot, previsão meteorológica, alertas e recomendações técnicas. Esta função melhora a acessibilidade para utilizadores de campo com dificuldade de leitura, mantendo sempre a resposta escrita visível no ecrã.")

    h1(doc, "6. Motor de Regras e Recomendações")
    p(doc, "O sistema calcula risco agronómico simples usando histórico climático, pragas/doenças activas e alertas não lidos. A pontuação é convertida em níveis como Baixo, Médio, Alto ou Crítico. Com base nisso, a aplicação sugere prioridade e texto preliminar de recomendação.")
    table(doc, ["Elemento", "Como funciona"], [
        ["Clima", "Temperatura elevada, humidade baixa ou vento forte aumentam o risco."],
        ["Pragas/Doenças", "Ocorrências activas aumentam a pontuação conforme a severidade."],
        ["Alertas", "Alertas activos também aumentam o risco do talhão/propriedade."],
        ["IA/Fallback", "Quando Gemini está activo, a recomendação é enriquecida pela IA; sem internet ou sem chave, o texto vem das regras internas."],
        ["Resultado", "O sistema sugere uma recomendação preliminar que deve ser revista pelo consultor."],
    ])

    h1(doc, "7. Filtros e Notificações")
    p(doc, "As listas principais possuem filtros para facilitar a consulta durante a operação e a defesa. É possível filtrar utilizadores por perfil e estado de solicitação, propriedades por nome/localização, recomendações por prioridade/estado, visitas por tipo, pragas por severidade/estado, alertas por tipo/severidade/estado e registos climáticos por propriedade ou descrição.")
    p(doc, "O painel também apresenta notificações no topo. Para o administrador, aparecem solicitações de perfil pendentes. Para os utilizadores com propriedades ou acesso técnico, aparecem alertas não lidos relacionados ao seu âmbito de acesso.")

    h1(doc, "8. Acessibilidade e Resposta por Voz")
    p(doc, "A resposta por voz foi adicionada para aproximar o sistema da realidade do campo. Muitos utilizadores podem ter dificuldade em ler textos longos no telemóvel ou computador, por isso a aplicação passa a oferecer leitura automática ou manual dos conteúdos mais importantes.")
    table(doc, ["Área", "Funcionamento por voz"], [
        ["Chatbot", "Ao entrar, reproduz uma mensagem de boas-vindas. Depois mostra a resposta escrita e pode ler a resposta em voz alta. Cada mensagem do assistente tem botão para ouvir novamente."],
        ["Página interna", "O botão Ouvir página lê o conteúdo principal da página actual."],
        ["Meteorologia", "O botão Ouvir previsão lê temperatura, humidade, vento e condição climática."],
        ["Alertas", "O botão Ouvir alertas lê os alertas activos ou informa quando não existem alertas."],
        ["Recomendações", "O botão Ouvir recomendação lê talhão, propriedade, prioridade, dados de solo, clima e recomendação técnica."],
    ])
    p(doc, "A função usa a síntese de voz do navegador. Por isso, não depende de servidor externo para transformar texto em fala. Caso o navegador não suporte leitura por voz, a aplicação mantém o texto escrito normalmente.")

    h1(doc, "9. Como Demonstrar na Defesa")
    number(doc, "Abrir o portal público e apresentar o objectivo da plataforma.")
    number(doc, "Entrar com utilizador administrador ou consultor.")
    number(doc, "Mostrar dashboard por perfil.")
    number(doc, "Abrir propriedades e talhões.")
    number(doc, "Criar ou visualizar uma visita técnica.")
    number(doc, "Mostrar uma recomendação e a exportação em PDF.")
    number(doc, "Abrir meteorologia, histórico climático e alertas.")
    number(doc, "Clicar em Ouvir página, Ouvir previsão, Ouvir alertas ou Ouvir recomendação para demonstrar acessibilidade.")
    number(doc, "Abrir Config. API e mostrar onde ficam a chave Gemini, o modelo e o interruptor de IA externa.")
    number(doc, "Explicar que, sem internet, o sistema continua a responder com regras internas e perguntas programadas.")

    h1(doc, "10. Tecnologias Realmente Usadas")
    table(doc, ["Tecnologia", "Uso na aplicação"], [
        ["Python", "Linguagem principal do back-end."],
        ["Django 5.0", "Framework web, rotas, views, modelos, autenticação e admin."],
        ["MySQL", "Base de dados relacional usada no ambiente local."],
        ["mysqlclient", "Conector entre Django e MySQL."],
        ["HTML/CSS/JavaScript", "Interface, templates e interacções do frontend."],
        ["Web Speech API", "Leitura por voz no navegador através de SpeechSynthesis."],
        ["Pillow", "Suporte a upload e manipulação de imagens."],
        ["requests", "Consulta às APIs meteorológicas."],
        ["Google/Gemini API", "Integração opcional para chatbot e recomendações automáticas quando configurada."],
        ["ReportLab", "Geração de PDF das recomendações."],
        ["python-decouple", "Leitura de variáveis do ficheiro .env."],
        ["Laragon", "Ambiente local para MySQL/MariaDB."],
        ["Open-Meteo/OpenWeatherMap", "Provedores meteorológicos suportados."],
    ])

    h1(doc, "11. Evoluções Futuras")
    bullet(doc, "Aprimorar prompts e auditoria das respostas geradas pela Gemini API.")
    bullet(doc, "Permitir escolha de voz, idioma e velocidade de leitura no perfil do utilizador.")
    bullet(doc, "Notificações por email, SMS ou WhatsApp.")
    bullet(doc, "Aplicação móvel.")
    bullet(doc, "Dashboards com gráficos avançados.")
    bullet(doc, "Deploy em servidor com HTTPS.")
    bullet(doc, "Validação com utilizadores reais da Hispatec.")

    footer(doc, "Relatório de Funcionamento - AgroVision / Hispatec")
    doc.save(OUT_FUNC)


def build_technical_report():
    doc = setup_doc(
        "Relatório Técnico Simplificado da Programação do AgroVision",
        "Guia para entender o código e fazer alterações básicas com segurança",
    )

    h1(doc, "1. Para Quem é Este Documento")
    p(doc, "Este guia foi escrito para uma pessoa que não tem muitos conhecimentos de programação, mas precisa entender como a aplicação AgroVision está organizada e onde deve mexer caso queira fazer pequenas alterações no futuro.")
    p(doc, "A regra principal é: antes de alterar qualquer ficheiro, fazer uma cópia do projecto ou usar Git. Alterações em modelos e base de dados exigem mais cuidado do que alterações em textos, imagens ou estilos.")

    h1(doc, "2. Como o Django Organiza a Aplicação")
    p(doc, "Django separa o projecto em módulos chamados apps. Cada app cuida de uma parte do sistema. Dentro de cada app existem ficheiros com responsabilidades diferentes.")
    table(doc, ["Ficheiro/Pasta", "Para que serve"], [
        ["models.py", "Define as tabelas da base de dados em formato de classes Python."],
        ["views.py", "Contém a lógica que recebe pedidos do navegador e devolve páginas ou respostas."],
        ["urls.py", "Define os endereços/rotas que o utilizador acessa."],
        ["forms.py", "Define formulários, validações e campos exibidos ao utilizador."],
        ["admin.py", "Configura como os dados aparecem no painel administrativo do Django."],
        ["templates/", "Contém páginas HTML mostradas no navegador."],
        ["static/", "Contém CSS, JavaScript, imagens e ficheiros visuais."],
        ["migrations/", "Guarda o histórico de alterações da base de dados."],
    ])

    h1(doc, "3. Apps do Projecto")
    table(doc, ["App", "Responsabilidade"], [
        ["contas", "Registo, login, perfis de utilizador, solicitação de perfil e permissões."],
        ["publico", "Páginas públicas: início, serviços, sobre e contacto."],
        ["dashboard", "Painéis por perfil e chatbot interno."],
        ["propriedades", "Propriedades, culturas e talhões."],
        ["consultoria", "Recomendações, visitas técnicas, fotos e pragas/doenças."],
        ["meteorologia", "Previsão meteorológica, registos climáticos e alertas."],
        ["config_sistema", "Configuração das APIs meteorológica e Google/Gemini."],
        ["agrovision", "Configurações principais do projecto, URLs globais e arranque."],
    ])

    h1(doc, "4. Caminhos Mais Importantes")
    table(doc, ["O que quer alterar", "Onde mexer primeiro"], [
        ["Textos da página inicial", "templates/publico/home.html"],
        ["Página de serviços", "templates/publico/servicos.html"],
        ["Página sobre", "templates/publico/sobre.html"],
        ["Página contacto", "templates/publico/contacto.html"],
        ["Layout geral do dashboard", "templates/base_dashboard.html"],
        ["Estilos visuais", "static/css/style.css"],
        ["Chatbot visual", "static/js/chatbot.js"],
        ["Leitura por voz", "static/js/voice-tools.js e botões data-speech nos templates."],
        ["Respostas do chatbot", "dashboard/views.py, config_sistema/ia.py ou painel admin em Respostas do Chatbot"],
        ["Regras de risco/recomendação", "consultoria/views.py"],
        ["Chave Google/Gemini", "Painel Config. API ou modelo ConfiguracaoAPI em config_sistema/models.py"],
        ["Notificações do painel", "dashboard/context_processors.py e templates/base_dashboard.html"],
        ["Filtros das listas", "Views de cada app e templates das listas."],
        ["Campos de utilizador", "contas/models.py"],
        ["Campos de propriedade/talhão", "propriedades/models.py"],
        ["Campos de recomendação/visita/praga", "consultoria/models.py"],
        ["Configuração da base de dados", "agrovision/settings.py e ficheiro .env"],
    ])

    h1(doc, "5. Fluxo de uma Página")
    number(doc, "O utilizador clica num link no navegador.")
    number(doc, "O ficheiro urls.py encontra a rota correspondente.")
    number(doc, "A view em views.py executa a lógica necessária.")
    number(doc, "A view consulta ou grava dados através dos models.py.")
    number(doc, "A view envia dados para um template HTML.")
    number(doc, "O navegador mostra a página com CSS e JavaScript.")

    h1(doc, "6. Como Alterar Textos e Imagens")
    p(doc, "Alterar textos e imagens é a mudança mais simples. Normalmente acontece em ficheiros HTML dentro da pasta templates ou em imagens dentro de static/img.")
    number(doc, "Abrir o ficheiro HTML correspondente.")
    number(doc, "Procurar o texto actual.")
    number(doc, "Substituir pelo novo texto.")
    number(doc, "Guardar o ficheiro.")
    number(doc, "Recarregar a página no navegador.")
    p(doc, "Para imagens, colocar o ficheiro em static/img e trocar o nome no atributo src do HTML. Evitar nomes com acentos ou espaços.")

    h1(doc, "7. Como Alterar Cores, Espaçamentos e Tamanhos")
    p(doc, "A maioria das alterações visuais deve ser feita no ficheiro static/css/style.css. Este ficheiro controla cores, botões, cartões, tabelas, sidebar e páginas públicas.")
    bullet(doc, "Para mudar cor, procurar códigos como #1f7033 ou variáveis CSS.")
    bullet(doc, "Para mudar tamanho de texto, procurar font-size.")
    bullet(doc, "Para mudar espaçamento, procurar padding, margin ou gap.")
    bullet(doc, "Para mudar largura, procurar width, max-width ou grid-template-columns.")

    h1(doc, "8. Como Alterar Regras de Recomendação")
    p(doc, "As recomendações preliminares são calculadas em consultoria/views.py. O sistema tenta usar Gemini quando a IA externa está activa; se a API não responder, usa regras simples locais, principalmente nas funções _risco_talhao e _sugestao_recomendacao.")
    table(doc, ["Função", "O que faz"], [
        ["_risco_talhao", "Calcula a pontuação de risco usando clima, pragas e alertas."],
        ["_sugestao_recomendacao", "Transforma o risco calculado em texto de recomendação preliminar."],
        ["_sugestao_recomendacao_automatica", "Tenta gerar texto com Gemini e, se falhar, usa o texto por regras."],
        ["nova_recomendacao", "Mostra o formulário e grava a recomendação final."],
    ])
    p(doc, "Se quiser mudar os limites do risco, deve alterar os valores usados para Baixo, Médio, Alto e Crítico. Depois é preciso testar uma recomendação para confirmar que o texto aparece correctamente.")

    h1(doc, "9. Como Funciona a Meteorologia")
    p(doc, "A meteorologia está em meteorologia/views.py. A função obter_previsao_api consulta OpenWeatherMap quando há chave configurada, ou Open-Meteo quando não há chave. A lista LOCALIDADES_AGRICOLAS_ANGOLA alimenta sugestões de províncias e municípios agrícolas no formulário. A função _guardar_registo_e_alerta cria registos climáticos e alertas simples quando identifica calor extremo ou vento forte.")
    table(doc, ["Parte", "Descrição"], [
        ["ConfiguracaoAPI", "Modelo que guarda provedor, chave, URL base e cidade padrão."],
        ["obter_previsao_api", "Consulta a API meteorológica externa."],
        ["RegistroClima", "Tabela que guarda temperatura, humidade, vento e descrição."],
        ["Alerta", "Tabela que guarda mensagens de risco para propriedades."],
    ])

    h1(doc, "10. Como Funciona o Chatbot e a IA")
    p(doc, "O chatbot funciona em modo híbrido. O endpoint fica em dashboard/chatbot/. O JavaScript em static/js/chatbot.js cria a janela flutuante, envia perguntas ao servidor, mostra as respostas e activa a leitura por voz quando o utilizador liga esta opção.")
    p(doc, "Quando a opção IA externa está activa no painel Config. API, a view chama config_sistema/ia.py, que envia um prompt para a API Google/Gemini com contexto limitado ao perfil do utilizador. Se não houver chave, internet ou resposta válida, a view usa respostas editáveis e regras programadas.")
    table(doc, ["Ficheiro", "Função"], [
        ["config_sistema/models.py", "Guarda ia_api_key, ia_modelo, ia_url_base e ia_ativo."],
        ["config_sistema/ia.py", "Executa a chamada HTTP para Gemini e devolve None em caso de falha."],
        ["dashboard/views.py", "Escolhe entre resposta Gemini, resposta editável e resposta programada."],
        ["consultoria/views.py", "Usa a mesma camada de IA para enriquecer recomendações automáticas."],
    ])

    h1(doc, "11. Como Funciona a Leitura por Voz")
    p(doc, "A leitura por voz foi implementada no frontend, usando a Web Speech API do navegador. Isso significa que o servidor Django não precisa gerar ficheiros de áudio. O texto continua escrito na página, e o navegador lê esse texto quando o utilizador clica no botão de voz.")
    table(doc, ["Ficheiro/Elemento", "Função"], [
        ["static/js/voice-tools.js", "Cria a função global AgroVisionVoice, lê textos com SpeechSynthesis e trata botões data-speech ou data-speech-target."],
        ["static/js/chatbot.js", "Reproduz a mensagem de boas-vindas ao entrar, lê respostas do assistente quando a voz está ligada e permite ouvir cada resposta novamente."],
        ["templates/base_dashboard.html", "Carrega voice-tools.js e mostra botões Ouvir página e Parar leitura."],
        ["templates/meteorologia/previsao.html", "Inclui botão Ouvir previsão com resumo climático."],
        ["templates/meteorologia/lista_alertas.html", "Inclui botão Ouvir alertas."],
        ["templates/consultoria/detalhe_recomendacao.html", "Inclui botão Ouvir recomendação."],
        ["static/css/style.css", "Define aparência dos botões de voz e do botão de ouvir no chatbot."],
    ])
    p(doc, "Para adicionar voz numa nova página, basta criar um botão com data-speech para texto directo ou data-speech-target para ler o conteúdo de um elemento da página. Exemplo: um botão com data-speech-target='.main-content' lê o conteúdo principal.")
    p(doc, "Limitação técnica: a voz depende do navegador e das vozes instaladas no sistema operativo. Em Chrome ou Edge costuma funcionar bem, mas a qualidade da voz pode variar conforme o computador.")

    h1(doc, "12. Notificações e Filtros")
    p(doc, "As notificações globais são geradas em dashboard/context_processors.py. Esse ficheiro calcula solicitações pendentes para o administrador e alertas não lidos para cada perfil. O template base_dashboard.html mostra o contador no topo do painel.")
    p(doc, "Os filtros ficam nas views e templates das listas. Eles usam parâmetros GET como q, tipo, status, prioridade, severidade e estado. Isto permite pesquisar sem alterar a base de dados.")

    h1(doc, "13. Alterações na Base de Dados")
    p(doc, "Mudanças nos ficheiros models.py podem alterar a base de dados. Depois de alterar campos, normalmente é necessário executar comandos de migração.")
    table(doc, ["Comando", "Função"], [
        ["python manage.py makemigrations", "Cria ficheiros de migração com as mudanças dos modelos."],
        ["python manage.py migrate", "Aplica as mudanças na base de dados."],
        ["python manage.py check", "Verifica problemas básicos no projecto."],
        ["python manage.py runserver", "Liga o servidor local."],
    ])
    p(doc, "Antes de mexer em models.py, fazer cópia da base de dados ou confirmar que os dados podem ser recriados.")

    h1(doc, "14. Cuidados Antes de Alterar")
    bullet(doc, "Não apagar migrations antigas sem saber o impacto.")
    bullet(doc, "Não colocar senhas ou chaves API directamente no código.")
    bullet(doc, "Não mudar nomes de campos no modelo sem actualizar formulários, views, templates e banco de dados.")
    bullet(doc, "Testar login, dashboard, propriedades, recomendações, meteorologia e alertas depois de qualquer mudança importante.")
    bullet(doc, "Guardar sempre uma cópia funcional antes de experimentar alterações grandes.")

    h1(doc, "15. Guia Rápido de Problemas Comuns")
    table(doc, ["Problema", "Onde verificar"], [
        ["Página não abre", "urls.py da app e agrovision/urls.py."],
        ["Erro ao guardar formulário", "forms.py, views.py e campos do model."],
        ["Imagem não aparece", "Caminho em static/img ou media/."],
        ["Dados não aparecem para utilizador", "Filtros por perfil nas funções _propriedades_visiveis."],
        ["PDF não gera", "consultoria/views.py e dependência ReportLab."],
        ["API meteorológica falha", "config_sistema, meteorologia/views.py, internet e chave API."],
        ["Admin de clima/alertas dá erro de data", "Verificar meteorologia/admin.py. O date_hierarchy foi removido para não depender das tabelas de timezone do MySQL local."],
        ["Chatbot não responde", "dashboard/views.py, static/js/chatbot.js e URL dashboard/chatbot/."],
        ["Voz não funciona", "Verificar se o navegador suporta SpeechSynthesis, se o som do computador está activo, se static/js/voice-tools.js foi carregado e se o navegador permite áudio após entrada no sistema."],
    ])

    footer(doc, "Relatório Técnico Simplificado - AgroVision / Hispatec")
    doc.save(OUT_TECH)


def validate(path):
    with zipfile.ZipFile(path) as z:
        required = {"word/document.xml", "word/styles.xml", "[Content_Types].xml"}
        missing = required - set(z.namelist())
        if missing:
            raise RuntimeError(f"DOCX incompleto: {path} {missing}")


if __name__ == "__main__":
    build_functional_report()
    build_technical_report()
    validate(OUT_FUNC)
    validate(OUT_TECH)
    print(OUT_FUNC.resolve())
    print(OUT_TECH.resolve())
